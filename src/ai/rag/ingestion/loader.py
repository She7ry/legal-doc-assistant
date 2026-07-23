"""文档解析与上传：PDF / DOCX / TXT → LangChain Document 列表。

``load_documents`` 按扩展名选 loader；PDF 可选 OCR（settings.pdf_ocr_enabled）。
``save_uploaded_file`` 将 API 上传的字节写入 ``data/uploads/``。
"""

from __future__ import annotations

import hashlib
import logging
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document

from ai.config.settings import settings

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".markdown", ".docx"}
INGEST_WARNINGS_METADATA_KEY = "ingest_warnings"


def file_sha256(path: Path) -> str:
    """计算文件 SHA-256，用作 file_id 与去重依据。"""
    digest = hashlib.sha256()
    with path.open("rb") as file:
        for block in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _safe_file_name(name: str) -> str:
    keep = []
    for char in name:
        if char.isalnum() or char in {".", "-", "_"}:
            keep.append(char)
        else:
            keep.append("_")
    return "".join(keep).strip("._") or "uploaded_document"


def save_uploaded_file(file_name: str, content: bytes, user_id: str | None = None) -> Path:
    """将 API 上传的原始字节写入 ``data/uploads/``（按用户分子目录）。"""
    safe_name = _safe_file_name(file_name)
    upload_dir = settings.upload_dir
    if user_id:
        upload_dir = upload_dir / _safe_file_name(user_id)

    upload_dir.mkdir(parents=True, exist_ok=True)
    target_path = upload_dir / f"{uuid4().hex}_{safe_name}"
    target_path.write_bytes(content)
    return target_path


def load_documents(path: Path) -> list[Document]:
    """按扩展名解析 PDF/DOCX/TXT，返回带 file_name、source metadata 的 Document 列表。"""
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持的文件类型：{suffix}")

    if suffix == ".pdf":
        documents = _load_pdf_documents(path)
    elif suffix == ".docx":
        documents = _load_docx_documents(path)
    else:
        documents = TextLoader(str(path), encoding="utf-8", autodetect_encoding=True).load()

    for document in documents:
        document.metadata["file_name"] = path.name
        document.metadata["source"] = str(path)
        document.metadata["file_extension"] = suffix

    return documents


def _load_pdf_documents(path: Path) -> list[Document]:
    """PyPDF 逐页解析；空页记录警告，可选 OCR 回填（settings.pdf_ocr_enabled）。"""
    documents = PyPDFLoader(str(path)).load()
    empty_pages = []

    for index, document in enumerate(documents):
        page = document.metadata.get("page", index)
        text = document.page_content or ""
        document.metadata["char_count"] = len(text)
        if not text.strip():
            empty_pages.append(int(page) if isinstance(page, int) else index)

    warnings = []
    if empty_pages:
        page_labels = ", ".join(str(page + 1) for page in empty_pages[:10])
        suffix = "..." if len(empty_pages) > 10 else ""
        warnings.append(
            f"PDF 第 {page_labels}{suffix} 页没有可提取文本，可能是扫描图片或包含不支持的版式。"
        )

    if empty_pages and settings.pdf_ocr_enabled:
        ocr_text_by_page, ocr_warnings = _ocr_pdf_pages(path, empty_pages)
        warnings.extend(ocr_warnings)
        for document in documents:
            page = document.metadata.get("page")
            if isinstance(page, int) and page in ocr_text_by_page:
                document.page_content = ocr_text_by_page[page]
                document.metadata["ocr_applied"] = True
                document.metadata["char_count"] = len(document.page_content)
    elif empty_pages:
        warnings.append(
            "OCR 当前未启用；如需索引扫描版 PDF，请安装 OCR 依赖并设置 "
            "DOC_ASSISTANT_PDF_OCR_ENABLED=true。"
        )

    if warnings and documents:
        _append_warnings(documents[0], warnings)

    return documents


def _ocr_pdf_pages(path: Path, pages: list[int]) -> tuple[dict[int, str], list[str]]:
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except ImportError:
        return {}, [
            "已请求 OCR，但尚未安装 pdf2image 和 pytesseract。"
        ]

    extracted: dict[int, str] = {}
    warnings = []
    for page in pages:
        try:
            images = convert_from_path(
                str(path),
                first_page=page + 1,
                last_page=page + 1,
            )
            if not images:
                warnings.append(f"OCR 未能生成 PDF 第 {page + 1} 页的图像。")
                continue
            text = pytesseract.image_to_string(images[0], lang=settings.pdf_ocr_lang).strip()
        except Exception:
            logger.warning("PDF OCR failed", extra={"path": str(path), "page": page}, exc_info=True)
            warnings.append(f"PDF 第 {page + 1} 页 OCR 失败。")
            continue

        if text:
            extracted[page] = text
        else:
            warnings.append(f"PDF 第 {page + 1} 页 OCR 未识别出文本。")

    return extracted, warnings


def _load_docx_documents(path: Path) -> list[Document]:
    doc = DocxDocument(str(path))
    documents: list[Document] = []
    current_heading: str | None = None
    current_parts: list[str] = []

    def flush_section() -> None:
        nonlocal current_parts
        content = "\n\n".join(part for part in current_parts if part).strip()
        if not content:
            return
        metadata = {"source": str(path), "file_name": path.name}
        if current_heading:
            metadata["section_heading"] = current_heading
        documents.append(Document(page_content=content, metadata=metadata))
        current_parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = getattr(paragraph.style, "name", "") or ""
        if style_name.casefold().startswith("heading"):
            flush_section()
            current_heading = text
            current_parts = [text]
        else:
            current_parts.append(text)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            current_parts.append("\n".join(rows))

    for section in doc.sections:
        header_text = "\n".join(paragraph.text.strip() for paragraph in section.header.paragraphs if paragraph.text.strip())
        footer_text = "\n".join(paragraph.text.strip() for paragraph in section.footer.paragraphs if paragraph.text.strip())
        if header_text:
            current_parts.append(f"Header:\n{header_text}")
        if footer_text:
            current_parts.append(f"Footer:\n{footer_text}")

    flush_section()
    if documents:
        return documents

    document = Document(page_content="", metadata={"source": str(path), "file_name": path.name})
    _append_warnings(document, ["DOCX file contained no extractable text."])
    return [document]


def _append_warnings(document: Document, warnings: list[str]) -> None:
    existing = document.metadata.get(INGEST_WARNINGS_METADATA_KEY)
    if isinstance(existing, list):
        document.metadata[INGEST_WARNINGS_METADATA_KEY] = existing + warnings
    else:
        document.metadata[INGEST_WARNINGS_METADATA_KEY] = warnings
