from __future__ import annotations

from types import SimpleNamespace

from docx import Document as DocxDocument

from ai.rag.ingestion import loader as document_loader


def test_save_uploaded_file_uses_user_directory_and_unique_names(tmp_path, monkeypatch) -> None:
    monkeypatch.setattr(document_loader, "settings", SimpleNamespace(upload_dir=tmp_path))

    first_path = document_loader.save_uploaded_file("Contract Copy.txt", b"one", user_id="acme")
    second_path = document_loader.save_uploaded_file("Contract Copy.txt", b"two", user_id="acme")

    assert first_path.parent == tmp_path / "acme"
    assert second_path.parent == tmp_path / "acme"
    assert first_path.name != second_path.name
    assert first_path.name.endswith("Contract_Copy.txt")
    assert first_path.read_bytes() == b"one"
    assert second_path.read_bytes() == b"two"


def test_markdown_extension_is_supported(tmp_path) -> None:
    path = tmp_path / "policy.markdown"
    path.write_text("# Policy\n\nUse written approval.", encoding="utf-8")

    documents = document_loader.load_documents(path)

    assert documents[0].page_content.startswith("# Policy")
    assert documents[0].metadata["file_extension"] == ".markdown"


def test_load_docx_extracts_paragraphs_and_tables(tmp_path) -> None:
    path = tmp_path / "contract.docx"
    source = DocxDocument()
    source.add_paragraph("Section 1 Term")
    table = source.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Party"
    table.cell(0, 1).text = "Acme"
    source.save(path)

    documents = document_loader.load_documents(path)

    assert "Section 1 Term" in documents[0].page_content
    assert "Party | Acme" in documents[0].page_content
    assert documents[0].metadata["file_extension"] == ".docx"


def test_load_docx_extracts_headers_and_footers(tmp_path) -> None:
    path = tmp_path / "contract.docx"
    source = DocxDocument()
    source.add_paragraph("Main agreement text.")
    source.sections[0].header.paragraphs[0].text = "Confidential header"
    source.sections[0].footer.paragraphs[0].text = "Page footer"
    source.save(path)

    documents = document_loader.load_documents(path)
    content = documents[0].page_content

    assert "Main agreement text." in content
    assert "Confidential header" in content
    assert "Page footer" in content
